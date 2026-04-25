"""
app/core/services/report_service.py
────────────────────────────────────
PDF (ReportLab) and Word (python-docx) report generation.
All reports are saved to ./reports/ and every generation
is audit-logged under Actions.REPORT_GENERATED.

Reports available:
  - generate_loan_agreement      ← loan agreement with collateral images
  - portfolio_summary_pdf / portfolio_summary_word
  - overdue_report_pdf
  - repayment_history_pdf
  - client_register_pdf / client_register_word
"""

import os
from datetime import date
from decimal import Decimal

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle,
    Paragraph, Spacer, HRFlowable,
    Image as RLImage,
)
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.services.audit_service import AuditService, Actions

REPORTS_DIR = "./reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

# ── Colour palette ─────────────────────────────────────────────────────────────
GREEN      = colors.HexColor("#1A5C1E")
GOLD       = colors.HexColor("#D4A017")
LIGHT_GREY = colors.HexColor("#F0F7F0")
NAVY       = colors.HexColor("#0D1B2A")

# Supported image extensions that can be rendered inline in PDF
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff"}


class ReportService:

    # ── PDF helpers ────────────────────────────────────────────────────────────

    @staticmethod
    @staticmethod
    def _base_pdf(filename: str, save_path: str = None):
        path = save_path if save_path else os.path.join(REPORTS_DIR, filename)
        doc  = SimpleDocTemplate(
            path, pagesize=A4,
            topMargin=2*cm, bottomMargin=2*cm,
            leftMargin=2*cm, rightMargin=2*cm,
        )
        styles = getSampleStyleSheet()
        return doc, styles, path

    @staticmethod
    def _header_elements(title: str) -> list:
        h1 = ParagraphStyle(
            "H1", fontSize=18, fontName="Helvetica-Bold",
            textColor=GREEN, spaceAfter=4)
        sub = ParagraphStyle(
            "Sub", fontSize=10, textColor=colors.grey, spaceAfter=12)
        return [
            Paragraph("BINGONGOLD CREDIT", h1),
            Paragraph(
                "together as one  |  Ham Tower, 4th Floor, Wandegeya, Kampala", sub),
            Paragraph(f"{title}  —  Generated: {date.today()}", sub),
            HRFlowable(width="100%", thickness=2, color=GREEN, spaceAfter=12),
        ]

    @staticmethod
    def _table_style(header_rows: int = 1) -> TableStyle:
        return TableStyle([
            ("BACKGROUND",     (0, 0), (-1, header_rows - 1), GREEN),
            ("TEXTCOLOR",      (0, 0), (-1, header_rows - 1), colors.white),
            ("FONTNAME",       (0, 0), (-1, header_rows - 1), "Helvetica-Bold"),
            ("FONTSIZE",       (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, header_rows), (-1, -1), [LIGHT_GREY, colors.white]),
            ("GRID",           (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING",    (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",   (0, 0), (-1, -1), 6),
            ("TOPPADDING",     (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",  (0, 0), (-1, -1), 4),
        ])

    @staticmethod
    def _footer_elements(row_count: int, report_name: str) -> list:
        foot = ParagraphStyle(
            "foot", fontSize=8, textColor=colors.grey,
            alignment=1, spaceBefore=8)
        return [
            Spacer(1, 0.4*cm),
            HRFlowable(width="100%", thickness=1, color=GREEN),
            Paragraph(
                f"{report_name}  |  Total records: {row_count}  |  "
                f"Bingongold Credit  |  {date.today()}",
                foot,
            ),
        ]

    # ── Word helpers ───────────────────────────────────────────────────────────

    @staticmethod
    def _base_word(title: str) -> Document:
        doc = Document()
        h   = doc.add_heading("BINGONGOLD CREDIT", 0)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x1E)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(
            "together as one  |  Ham Tower, 4th Floor, Wandegeya, Kampala")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size      = Pt(9)
        sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_heading(title, 1)
        doc.add_paragraph(f"Generated: {date.today()}")
        doc.add_paragraph()
        return doc

    @staticmethod
    def _word_table_header(table, headers: list):
        row = table.rows[0].cells
        for i, h in enumerate(headers):
            row[i].text = h
            run = row[i].paragraphs[0].runs[0]
            run.bold                = True
            run.font.color.rgb      = RGBColor(0x1A, 0x5C, 0x1E)

    # ── Internal: render one collateral item into PDF elements ─────────────────

    @staticmethod
    def _collateral_elements(coll_data: list, term_s, muted_s) -> list:
        """
        Build a list of ReportLab flowables for all collateral documents.

        Args:
            coll_data:  List of (description, file_path, file_name) tuples.
            term_s:     ParagraphStyle for body text.
            muted_s:    ParagraphStyle for muted/caption text.
        """
        elements = []

        if not coll_data:
            elements.append(Paragraph(
                "No collateral documents were attached to this loan.", muted_s))
            return elements

        for idx, (desc, fpath, fname) in enumerate(coll_data, 1):
            ext = os.path.splitext(fpath)[1].lower()

            # Document label
            elements.append(Paragraph(
                f"<b>Document {idx}:</b>  {desc}  <font color='grey'>({fname})</font>",
                term_s,
            ))

            if ext in IMAGE_EXTS:
                if os.path.exists(fpath):
                    try:
                        # Read actual image dimensions and scale to fit page
                        reader  = ImageReader(fpath)
                        iw, ih  = reader.getSize()
                        max_w   = 14 * cm
                        max_h   = 9  * cm
                        ratio   = min(max_w / iw, max_h / ih, 1.0)
                        draw_w  = iw * ratio
                        draw_h  = ih * ratio

                        img = RLImage(fpath, width=draw_w, height=draw_h)
                        elements.append(Spacer(1, 0.2*cm))
                        elements.append(img)
                        elements.append(Paragraph(
                            f"{desc} — {fname}", muted_s))
                    except Exception as e:
                        elements.append(Paragraph(
                            f"[Image could not be rendered: {e}]", muted_s))
                else:
                    elements.append(Paragraph(
                        f"[Image file not found on disk: {fpath}]", muted_s))

            elif ext == ".pdf":
                elements.append(Paragraph(
                    f"[PDF document — printed separately: {fname}]", muted_s))

            else:
                elements.append(Paragraph(
                    f"[Document type {ext or 'unknown'}: {fname}]", muted_s))

            elements.append(Spacer(1, 0.5*cm))

        return elements

    # ── Loan Agreement ─────────────────────────────────────────────────────────

    @staticmethod
    def generate_loan_agreement(loan, client, generated_by_id: int = None) -> str:
        """
        Generate a printable PDF loan agreement including:
        - Borrower details
        - Full loan financial breakdown
        - Terms and conditions
        - Signature block
        - All collateral documents (images rendered inline)

        Args:
            loan:             Loan model instance.
            client:           Client model instance.
            generated_by_id:  ID of the user printing this (for audit log).
        """
        doc, _, path = ReportService._base_pdf(
            f"loan_agreement_{loan.loan_number}.pdf")

        elements = ReportService._header_elements("LOAN AGREEMENT")

        # ── Paragraph styles ───────────────────────────────────────────────
        intro_s = ParagraphStyle(
            "intro", fontSize=10, textColor=colors.black,
            spaceAfter=8, leading=16)
        section_s = ParagraphStyle(
            "sec", fontSize=11, fontName="Helvetica-Bold",
            textColor=GREEN, spaceAfter=6, spaceBefore=4)
        term_s = ParagraphStyle(
            "ts", fontSize=9, textColor=colors.black,
            spaceAfter=4, leading=14)
        muted_s = ParagraphStyle(
            "mt", fontSize=8, textColor=colors.grey,
            spaceAfter=2, leading=12)

        # ── Intro paragraph ────────────────────────────────────────────────
        elements.append(Paragraph(
            f"This Loan Agreement is entered into on <b>{date.today()}</b> between "
            f"<b>Bingongold Credit</b>, Ham Tower, 4th Floor, Wandegeya, Kampala "
            f"(the <b>Lender</b>) and the borrower named below (the <b>Borrower</b>).",
            intro_s,
        ))
        elements.append(Spacer(1, 0.3*cm))

        # ── Borrower & loan details table ──────────────────────────────────
        client_name  = client.full_name    if client else "—"
        client_nin   = client.nin          if client else "—"
        client_phone = client.phone_number if client else "—"
        client_addr  = getattr(client, "address", None) or "—"

        tdata = [
            ["BORROWER DETAILS", ""],
            ["Full Name",     client_name],
            ["NIN",           client_nin],
            ["Phone",         client_phone],
            ["Address",       client_addr],
            ["", ""],
            ["LOAN DETAILS",  ""],
            ["Loan Number",   loan.loan_number],
            ["Loan Type",     loan.loan_type.value if loan.loan_type else "—"],
            ["Principal",     f"UGX {float(loan.principal_amount):,.0f}"],
            ["Interest (10%)",
             f"UGX {float(loan.total_interest):,.0f}"      if loan.total_interest      else "—"],
            ["Total Repayable",
             f"UGX {float(loan.total_repayable):,.0f}"     if loan.total_repayable     else "—"],
            ["Monthly Installment",
             f"UGX {float(loan.monthly_installment):,.0f}" if loan.monthly_installment else "—"],
            ["Duration",         f"{loan.duration_months} months"],
            ["Purpose",          loan.purpose or "—"],
            ["Application Date", str(loan.application_date) if loan.application_date else "—"],
        ]

        t = Table(tdata, colWidths=[6*cm, 11*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0),  (-1, 0),  GREEN),
            ("TEXTCOLOR",     (0, 0),  (-1, 0),  colors.white),
            ("FONTNAME",      (0, 0),  (-1, 0),  "Helvetica-Bold"),
            ("SPAN",          (0, 0),  (-1, 0)),
            ("BACKGROUND",    (0, 6),  (-1, 6),  GREEN),
            ("TEXTCOLOR",     (0, 6),  (-1, 6),  colors.white),
            ("FONTNAME",      (0, 6),  (-1, 6),  "Helvetica-Bold"),
            ("SPAN",          (0, 6),  (-1, 6)),
            ("FONTNAME",      (0, 1),  (0, 5),   "Helvetica-Bold"),
            ("FONTNAME",      (0, 7),  (0, -1),  "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0, 1),  (-1, 5),  [LIGHT_GREY, colors.white]),
            ("ROWBACKGROUNDS",(0, 7),  (-1, -1), [LIGHT_GREY, colors.white]),
            ("FONTSIZE",      (0, 0),  (-1, -1), 10),
            ("TOPPADDING",    (0, 0),  (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0),  (-1, -1), 5),
            ("LEFTPADDING",   (0, 0),  (-1, -1), 8),
            ("GRID",          (0, 0),  (-1, -1), 0.25, colors.lightgrey),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 0.8*cm))

        # ── Terms and conditions ───────────────────────────────────────────
        elements.append(Paragraph("TERMS AND CONDITIONS", section_s))
        for term in [
            "1. The Borrower agrees to repay the full amount as per the repayment schedule above.",
            "2. Interest is charged at a flat rate of 10% on the principal amount.",
            "3. Payments must be made on or before the agreed due date each month.",
            "4. Late payments may attract a penalty as determined by Bingongold Credit.",
            "5. The Borrower confirms that all information provided is true and accurate.",
            "6. Any collateral pledged remains under the Lender's charge until full repayment.",
            "7. The Lender reserves the right to pursue legal action in the event of default.",
            "8. This agreement is governed by the laws of the Republic of Uganda.",
            "9. Any disputes shall be resolved through courts of competent jurisdiction in Kampala.",
        ]:
            elements.append(Paragraph(term, term_s))

        elements.append(Spacer(1, 1*cm))

        # ── Signature block ────────────────────────────────────────────────
        sig_data = [
            [Paragraph("<b>Borrower Signature:</b>", term_s),
             Paragraph("", term_s),
             Paragraph("<b>Authorised Officer:</b>", term_s)],
            [Paragraph("_______________________", term_s),
             Paragraph("", term_s),
             Paragraph("_______________________", term_s)],
            [Paragraph(client_name, term_s),
             Paragraph("", term_s),
             Paragraph("Bingongold Credit", term_s)],
            [Paragraph("Date: _______________", term_s),
             Paragraph("", term_s),
             Paragraph(f"Date: {date.today()}", term_s)],
        ]
        sig_t = Table(sig_data, colWidths=[7*cm, 3*cm, 7*cm])
        sig_t.setStyle(TableStyle([
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("ALIGN",      (2, 0), (2, -1),  "RIGHT"),
        ]))
        elements.append(sig_t)
        elements.append(Spacer(1, 0.5*cm))
        elements.append(Paragraph(
            "This document is computer-generated and is valid without a physical stamp "
            "when signed by both parties.",
            muted_s,
        ))

        # ── Collateral documents ───────────────────────────────────────────
        elements.append(Spacer(1, 0.6*cm))
        elements.append(HRFlowable(width="100%", thickness=1.5, color=GREEN, spaceAfter=8))
        elements.append(Paragraph("COLLATERAL DOCUMENTS", section_s))
        elements.append(Paragraph(
            "The following documents have been submitted as collateral security for this loan.",
            muted_s,
        ))
        elements.append(Spacer(1, 0.3*cm))

        # Load collateral records from DB
        try:
            from app.database.connection import get_db
            from app.core.models.collateral import Collateral

            with get_db() as db:
                collaterals = (
                    db.query(Collateral)
                    .filter_by(loan_id=loan.id)
                    .order_by(Collateral.created_at)
                    .all()
                )
                coll_data = [
                    (c.description, c.file_path, c.file_name)
                    for c in collaterals
                ]
        except Exception as e:
            coll_data = []
            elements.append(Paragraph(
                f"[Could not load collateral records: {e}]", muted_s))

        elements.extend(
            ReportService._collateral_elements(coll_data, term_s, muted_s))

        elements.extend(ReportService._footer_elements(1, "Loan Agreement"))

        doc.build(elements)

        AuditService.log(
            action      = Actions.REPORT_GENERATED,
            user_id     = generated_by_id,
            entity_type = "Report",
            description = (
                f"Loan Agreement generated: {loan.loan_number} "
                f"| Client: {client_name} "
                f"| Collateral items: {len(coll_data)}"
            ),
        )
        return path

    # ── Portfolio Summary ──────────────────────────────────────────────────────

    @staticmethod
    def portfolio_summary_pdf(generated_by_id: int = None) -> str:
        """Generate a full portfolio summary PDF report."""
        from app.core.services.loan_service import LoanService
        from app.core.services.client_service import ClientService

        doc, _, path = ReportService._base_pdf(
            f"portfolio_summary_{date.today()}.pdf")
        elements = ReportService._header_elements("Portfolio Summary Report")

        counts  = LoanService.count_by_status()
        total   = LoanService.total_portfolio_value()
        clients = ClientService.count_clients()

        summary_data = [
            ["Metric",                   "Value"],
            ["Total Registered Clients", str(clients)],
            ["Total Active Portfolio",   f"UGX {float(total):,.0f}"],
            ["Pending Loans",            str(counts.get("pending",   0))],
            ["Approved Loans",           str(counts.get("approved",  0))],
            ["Active Loans",             str(counts.get("active",    0))],
            ["Completed Loans",          str(counts.get("completed", 0))],
            ["Defaulted Loans",          str(counts.get("defaulted", 0))],
        ]
        t = Table(summary_data, colWidths=[10*cm, 7*cm])
        t.setStyle(ReportService._table_style())
        elements.extend([t, Spacer(1, 0.6*cm)])

        loans     = LoanService.get_all_loans()
        loan_data = [["Loan No.", "Client", "Type", "Principal (UGX)", "Status", "Due Date"]]
        for loan in loans:
            client = ClientService.get_client_by_id(loan.client_id)
            loan_data.append([
                loan.loan_number,
                client.full_name if client else "—",
                loan.loan_type.value if loan.loan_type else "—",
                f"{float(loan.principal_amount):,.0f}",
                loan.status.value.upper(),
                str(loan.due_date) if loan.due_date else "—",
            ])

        t2 = Table(loan_data, colWidths=[3*cm, 4.5*cm, 3.5*cm, 3*cm, 2.5*cm, 2.5*cm])
        t2.setStyle(ReportService._table_style())
        elements.append(t2)
        elements.extend(ReportService._footer_elements(len(loans), "Portfolio Summary"))

        doc.build(elements)

        AuditService.log(
            action      = Actions.REPORT_GENERATED,
            user_id     = generated_by_id,
            entity_type = "Report",
            description = (
                f"Portfolio Summary PDF generated | "
                f"Loans: {len(loans)} | Active portfolio: UGX {float(total):,.0f}"
            ),
        )
        return path

    @staticmethod
    def portfolio_summary_word(generated_by_id: int = None) -> str:
        """Generate a full portfolio summary Word report."""
        from app.core.services.loan_service import LoanService
        from app.core.services.client_service import ClientService

        path = os.path.join(REPORTS_DIR, f"portfolio_summary_{date.today()}.docx")
        doc  = ReportService._base_word("Portfolio Summary Report")

        counts  = LoanService.count_by_status()
        total   = LoanService.total_portfolio_value()
        clients = ClientService.count_clients()

        doc.add_heading("Summary Statistics", 2)
        stats = [
            ("Total Clients",   str(clients)),
            ("Total Portfolio", f"UGX {float(total):,.0f}"),
            ("Active Loans",    str(counts.get("active",    0))),
            ("Completed Loans", str(counts.get("completed", 0))),
            ("Defaulted Loans", str(counts.get("defaulted", 0))),
        ]
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        ReportService._word_table_header(tbl, ["Metric", "Value"])
        for label, val in stats:
            row = tbl.add_row().cells
            row[0].text = label
            row[1].text = val

        doc.add_paragraph()
        doc.add_heading("Loan Portfolio Detail", 2)
        loans   = LoanService.get_all_loans()
        headers = ["Loan No.", "Client", "Type", "Principal (UGX)", "Status", "Due Date"]
        tbl2    = doc.add_table(rows=1, cols=len(headers))
        tbl2.style = "Table Grid"
        ReportService._word_table_header(tbl2, headers)
        for loan in loans:
            client = ClientService.get_client_by_id(loan.client_id)
            row    = tbl2.add_row().cells
            row[0].text = loan.loan_number
            row[1].text = client.full_name if client else "—"
            row[2].text = loan.loan_type.value if loan.loan_type else "—"
            row[3].text = f"{float(loan.principal_amount):,.0f}"
            row[4].text = loan.status.value.upper()
            row[5].text = str(loan.due_date) if loan.due_date else "—"

        doc.save(path)

        AuditService.log(
            action      = Actions.REPORT_GENERATED,
            user_id     = generated_by_id,
            entity_type = "Report",
            description = (
                f"Portfolio Summary Word report generated | "
                f"Loans: {len(loans)} | Active portfolio: UGX {float(total):,.0f}"
            ),
        )
        return path

    # ── Overdue Report ─────────────────────────────────────────────────────────

    @staticmethod
    def overdue_report_pdf(generated_by_id: int = None) -> str:
        """Generate a PDF of all overdue loans."""
        from app.core.services.loan_service import LoanService
        from app.core.services.client_service import ClientService

        doc, styles, path = ReportService._base_pdf(
            f"overdue_loans_{date.today()}.pdf")
        elements = ReportService._header_elements("Overdue Loans Report")

        overdue = LoanService.get_overdue_loans()
        if not overdue:
            elements.append(Paragraph(
                "No overdue loans as of today.", styles["Normal"]))
        else:
            data = [["Loan No.", "Client", "Phone", "Principal (UGX)", "Due Date", "Days Overdue"]]
            for loan in overdue:
                client = ClientService.get_client_by_id(loan.client_id)
                days   = (date.today() - loan.due_date).days if loan.due_date else 0
                data.append([
                    loan.loan_number,
                    client.full_name    if client else "—",
                    client.phone_number if client else "—",
                    f"UGX {float(loan.principal_amount):,.0f}",
                    str(loan.due_date),
                    str(days),
                ])
            t = Table(data, colWidths=[3*cm, 4*cm, 3*cm, 3.5*cm, 3*cm, 2.5*cm])
            t.setStyle(ReportService._table_style())
            elements.append(t)
            elements.extend(
                ReportService._footer_elements(len(overdue), "Overdue Loans"))

        doc.build(elements)

        AuditService.log(
            action      = Actions.REPORT_GENERATED,
            user_id     = generated_by_id,
            entity_type = "Report",
            description = f"Overdue Loans PDF generated | Overdue count: {len(overdue)}",
        )
        return path

    # ── Repayment History ──────────────────────────────────────────────────────

    @staticmethod
    def repayment_history_pdf(generated_by_id: int = None) -> str:
        """Generate a PDF of recent repayment history."""
        from app.core.services.repayment_service import RepaymentService
        from app.core.services.loan_service import LoanService

        doc, _, path = ReportService._base_pdf(
            f"repayment_history_{date.today()}.pdf")
        elements = ReportService._header_elements("Repayment History Report")

        repayments = RepaymentService.get_all_recent_repayments(limit=200)
        all_loans  = {l.id: l for l in LoanService.get_all_loans()}

        data = [["Receipt", "Loan No.", "Amount (UGX)", "Date", "Method"]]
        for r in repayments:
            loan = all_loans.get(r.loan_id)
            data.append([
                r.receipt_number,
                loan.loan_number if loan else "—",
                f"{float(r.amount):,.0f}",
                str(r.payment_date),
                r.payment_method.value if r.payment_method else "—",
            ])

        t = Table(data, colWidths=[4*cm, 3.5*cm, 3.5*cm, 3*cm, 3*cm])
        t.setStyle(ReportService._table_style())
        elements.append(t)
        elements.extend(
            ReportService._footer_elements(len(repayments), "Repayment History"))

        doc.build(elements)

        AuditService.log(
            action      = Actions.REPORT_GENERATED,
            user_id     = generated_by_id,
            entity_type = "Report",
            description = f"Repayment History PDF generated | Records: {len(repayments)}",
        )
        return path

    # ── Client Register ────────────────────────────────────────────────────────

    @staticmethod
    def client_register_pdf(generated_by_id: int = None) -> str:
        """Generate a PDF client register."""
        from app.core.services.client_service import ClientService

        doc, _, path = ReportService._base_pdf(
            f"client_register_{date.today()}.pdf")
        elements = ReportService._header_elements("Client Register")

        clients = ClientService.get_all_clients()
        data    = [["#", "Full Name", "NIN", "Phone", "District", "Occupation"]]
        for i, c in enumerate(clients, 1):
            data.append([
                str(i),
                c.full_name,
                c.nin          or "—",
                c.phone_number,
                getattr(c, "district",   None) or "—",
                getattr(c, "occupation", None) or "—",
            ])

        t = Table(data, colWidths=[1*cm, 5*cm, 3*cm, 3*cm, 3*cm, 3*cm])
        t.setStyle(ReportService._table_style())
        elements.append(t)
        elements.extend(
            ReportService._footer_elements(len(clients), "Client Register"))

        doc.build(elements)

        AuditService.log(
            action      = Actions.REPORT_GENERATED,
            user_id     = generated_by_id,
            entity_type = "Report",
            description = f"Client Register PDF generated | Clients: {len(clients)}",
        )
        return path

    @staticmethod
    def client_register_word(generated_by_id: int = None) -> str:
        """Generate a Word client register."""
        from app.core.services.client_service import ClientService

        path    = os.path.join(REPORTS_DIR, f"client_register_{date.today()}.docx")
        doc     = ReportService._base_word("Client Register")
        clients = ClientService.get_all_clients()

        headers = ["Full Name", "NIN", "Phone", "District", "Occupation"]
        tbl     = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        ReportService._word_table_header(tbl, headers)

        for c in clients:
            row    = tbl.add_row().cells
            row[0].text = c.full_name
            row[1].text = c.nin or "—"
            row[2].text = c.phone_number
            row[3].text = getattr(c, "district",   None) or "—"
            row[4].text = getattr(c, "occupation", None) or "—"

        doc.save(path)

        AuditService.log(
            action      = Actions.REPORT_GENERATED,
            user_id     = generated_by_id,
            entity_type = "Report",
            description = f"Client Register Word report generated | Clients: {len(clients)}",
        )
        return path